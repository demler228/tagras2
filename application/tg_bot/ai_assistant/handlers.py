import os
import re
import logging
import json
from glob import glob
from difflib import SequenceMatcher
from pathlib import Path
from datetime import datetime
import asyncio
from docx import Document as DocxDocument
from pypdf import PdfReader
from moviepy.editor import VideoFileClip
import whisper
from aiogram import Router, types, F
from aiogram.types import Message
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from gigachat import GigaChat
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from application.tg_bot.ai_assistant.keyboards.get_exit_button import get_exit_button_ai
from application.tg_bot.filters.is_admin import is_admin
from application.tg_bot.menu.personal_actions.keyboards.menu_keyboard import get_main_menu_keyboard
from aiogram.types.input_file import FSInputFile

logging.basicConfig(
    filename='assistant.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

auth = "Y2U3MmFkYTEtMGIzNC00M2UwLTliNGYtYjRhMGFhODUzYTFhOjRhZWMyODJkLWQ2NDktNDJiOS1hYjEwLTQ2ODUyNjAwZjBlYQ=="
giga = GigaChat(
    credentials=auth,
    model='GigaChat:latest',
    verify_ssl_certs=False
)

whisper_model = whisper.load_model("base")


class AIState(StatesGroup):
    active = State()


router = Router()

vectorstore = None
MAX_FILE_SIZE = 50 * 1024 * 1024


def similar(a, b):
    return SequenceMatcher(None, a.lower(), b.lower()).ratio()


def clean_text(text):
    text = re.sub(r'[^\w\s.,!?]', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\n+', '\n', text)
    return text.strip()


def process_pdf(pdf_path):
    try:
        if os.path.getsize(pdf_path) > MAX_FILE_SIZE:
            logging.warning(f"Файл {pdf_path} слишком большой, пропускаем")
            return ""
        text = ""
        reader = PdfReader(pdf_path)
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        return clean_text(text)
    except Exception as e:
        logging.error(f"Ошибка обработки PDF {pdf_path}: {e}")
        return ""


def process_docx(docx_path):
    try:
        if os.path.getsize(docx_path) > MAX_FILE_SIZE:
            logging.warning(f"Файл {docx_path} слишком большой, пропускаем")
            return ""
        doc = DocxDocument(docx_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        return clean_text(text)
    except Exception as e:
        logging.error(f"Ошибка обработки DOCX {docx_path}: {e}")
        return ""


def process_text(text_path):
    try:
        if os.path.getsize(text_path) > MAX_FILE_SIZE:
            logging.warning(f"Файл {text_path} слишком большой, пропускаем")
            return ""
        with open(text_path, "r", encoding="utf-8") as f:
            text = f.read()
        return clean_text(text)
    except Exception as e:
        logging.error(f"Ошибка обработки TXT {text_path}: {e}")
        return ""


async def process_video(video_path):
    cache_path = video_path.with_suffix(".txt")
    if cache_path.exists():
        try:
            with open(cache_path, "r", encoding="utf-8") as f:
                return clean_text(f.read())
        except Exception as e:
            logging.error(f"Ошибка чтения кэша {cache_path}: {e}")

    try:
        if os.path.getsize(video_path) > MAX_FILE_SIZE:
            logging.warning(f"Видео {video_path} слишком большое, пропускаем")
            return ""
        temp_audio = "temp_audio.wav"
        video = VideoFileClip(video_path)
        audio = video.audio
        audio.write_audiofile(temp_audio)
        result = await asyncio.to_thread(whisper_model.transcribe, temp_audio)
        os.remove(temp_audio)
        video.close()

        with open(cache_path, "w", encoding="utf-8") as f:
            f.write(result["text"])

        return clean_text(result["text"])
    except Exception as e:
        logging.error(f"Ошибка обработки видео {video_path}: {e}")
        return ""


def get_files_metadata(materials_path):
    files_metadata = {}
    file_extensions = ["*.pdf", "*.docx", "*.txt", "*.mp4", "*.avi", "*.mov", "*.mkv"]
    for ext in file_extensions:
        for file_path in materials_path.glob(ext):
            mtime = os.path.getmtime(file_path)
            files_metadata[str(file_path)] = mtime
    return files_metadata


def load_files_metadata():
    metadata_path = Path(__file__).parent / "files_metadata.json"
    if metadata_path.exists():
        try:
            with open(metadata_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            logging.error(f"Ошибка чтения метаданных файлов: {e}")
    return {}


def save_files_metadata(metadata):
    metadata_path = Path(__file__).parent / "files_metadata.json"
    try:
        with open(metadata_path, "w", encoding="utf-8") as f:
            json.dump(metadata, f, ensure_ascii=False)
    except Exception as e:
        logging.error(f"Ошибка сохранения метаданных файлов: {e}")


async def load_documents():
    docs = []
    materials_path = Path(__file__).parent / "materials"
    if not materials_path.exists():
        logging.error(f"Папка {materials_path} не найдена!")
        return docs

    # Обработка PDF
    for pdf_path in materials_path.glob("*.pdf"):
        try:
            logging.info(f"Обработка PDF: {pdf_path}")
            text = process_pdf(pdf_path)
            if text:
                docs.append(Document(
                    page_content=text,
                    metadata={"source": str(pdf_path), "type": "pdf"}
                ))
        except Exception as e:
            logging.error(f"Ошибка при обработке PDF {pdf_path}: {e}")

    # Обработка DOCX
    for docx_path in materials_path.glob("*.docx"):
        try:
            logging.info(f"Обработка DOCX: {docx_path}")
            text = process_docx(docx_path)
            if text:
                docs.append(Document(
                    page_content=text,
                    metadata={"source": str(docx_path), "type": "docx"}
                ))
        except Exception as e:
            logging.error(f"Ошибка при обработке DOCX {docx_path}: {e}")

    # Обработка TXT
    for txt_path in materials_path.glob("*.txt"):
        try:
            logging.info(f"Обработка TXT: {txt_path}")
            text = process_text(txt_path)
            if text:
                docs.append(Document(
                    page_content=text,
                    metadata={"source": str(txt_path), "type": "txt"}
                ))
        except Exception as e:
            logging.error(f"Ошибка при обработке TXT {txt_path}: {e}")

    # Обработка видео
    video_extensions = ['*.mp4', '*.avi', '*.mov', '*.mkv']
    for ext in video_extensions:
        for video_path in materials_path.glob(ext):
            try:
                logging.info(f"Обработка видео: {video_path}")
                text = await process_video(video_path)
                if text:
                    docs.append(Document(
                        page_content=f"[ВИДЕО {video_path.name}]:\n{text}",
                        metadata={"source": str(video_path), "type": "video"}
                    ))
            except Exception as e:
                logging.error(f"Ошибка при обработке видео {video_path}: {e}")

    logging.info(f"Загружено {len(docs)} файлов (документы, видео и текстовые файлы)")
    return docs


async def create_vector_store(docs):
    if not docs:
        logging.warning("Нет документов для обработки")
        return None

    index_path = Path(__file__).parent / "faiss_index"
    metadata_path = Path(__file__).parent / "files_metadata.json"
    materials_path = Path(__file__).parent / "materials"

    try:
        if index_path.exists() and metadata_path.exists():
            current_files = get_files_metadata(materials_path)
            saved_files = load_files_metadata()

            if current_files == saved_files:
                embeddings = HuggingFaceEmbeddings(model_name="cointegrated/LaBSE-en-ru")
                logging.info("Загрузка сохранённого векторного хранилища")
                return FAISS.load_local(str(index_path), embeddings, allow_dangerous_deserialization=True)

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=100,
            length_function=len,
        )
        splits = text_splitter.split_documents(docs)
        embeddings = HuggingFaceEmbeddings(model_name="cointegrated/LaBSE-en-ru")

        vectorstore = FAISS.from_documents(splits, embeddings)
        vectorstore.save_local(str(index_path))

        files_metadata = get_files_metadata(materials_path)
        save_files_metadata(files_metadata)

        logging.info("Векторное хранилище создано и сохранено")
        return vectorstore
    except Exception as e:
        logging.error(f"Ошибка при создании/загрузке векторного хранилища: {e}")
        return None


async def initialize_vector_store():
    global vectorstore
    docs = await load_documents()
    logging.info(f"Загружено документов: {len(docs)}")
    if not docs:
        logging.error("Документы не загружены!")
    vectorstore = await create_vector_store(docs)
    if vectorstore:
        logging.info("Векторное хранилище успешно инициализировано")
    else:
        logging.error("Не удалось создать векторное хранилище")


@router.callback_query(F.data == "ai_assistant_button")
async def handle_ai_assistant_button(callback_query: types.CallbackQuery, state: FSMContext):
    global vectorstore

    if vectorstore is None:
        msg = await callback_query.message.answer("Идёт инициализация базы знаний...")
        await initialize_vector_store()
        await msg.delete()
        if vectorstore is None:
            await callback_query.message.answer(
                "Не удалось загрузить базу знаний. Попробуйте позже.",
                reply_markup=get_exit_button_ai()
            )
            return

    await state.set_state(AIState.active)
    await callback_query.message.answer(
        "Привет! Я ваш виртуальный помощник в компании Таграс.\n"
        "Можете задать мне вопрос о документах, видеофайлах, текстовых файлах или просто пообщаться.",
        reply_markup=get_exit_button_ai()
    )


@router.message(F.text, AIState.active)
async def handle_text_message(message: Message, state: FSMContext):
    user_input = message.text

    if user_input.lower() in ["вернуться в меню", "выход"]:
        await message.answer("Режим общения с ИИ деактивирован.\n\nДобро пожаловать в меню",
                             reply_markup=get_main_menu_keyboard(is_admin(message.from_user.id)))
        await state.clear()
        return

    try:
        if not vectorstore:
            await message.answer(
                "База знаний не загружена. Попробуйте позже.",
                reply_markup=get_exit_button_ai()
            )
            return

        relevant_docs = vectorstore.similarity_search(user_input, k=5)
        context = "\n\n".join([
                                  f"[{doc.metadata.get('type', 'unknown').upper()} {Path(doc.metadata.get('source', '')).name}]:\n{doc.page_content}"
                                  for doc in relevant_docs])
        source_files = list(set([doc.metadata["source"] for doc in relevant_docs if "source" in doc.metadata]))

        system_prompt = (
            "Ты - ассистент компании Таграс. Отвечай на вопросы на русском языке, используя информацию из предоставленных документов, видео или текстовых файлов. "
            "Если информация отсутствует в контексте, используй свои знания, но отвечай только по темам, связанным с деятельностью компании Таграс. "
            "Укажи тип источника (PDF, Word, видео, текст) и название, если они использованы в ответе."
        )

        user_prompt = f"""
        Контекст:
        {context if context else "Информация в документах отсутствует."}

        Вопрос:
        {user_input}

        Ответь максимально точно на русском языке. Если информация взята из контекста, укажи тип и название источника. "
        "Если контекст пуст, ответь на основе своих знаний, указав, что информация не найдена в базе знаний.
        """

        response = await asyncio.to_thread(giga.chat, system_prompt + "\n" + user_prompt)
        answer = response.choices[0].message.content

        if not answer.strip():
            answer = "Информация по этому вопросу отсутствует в базе знаний. На основе моих знаний: ответа нет."

        await message.answer(answer, reply_markup=get_exit_button_ai())

        if source_files:
            await message.answer("Вот файлы, которые могут содержать полезную информацию:")

            for file_path in source_files[:3]:
                try:
                    if not os.path.exists(file_path):
                        continue
                    file_ext = file_path.split('.')[-1].lower()
                    file_name = os.path.basename(file_path)
                    file = FSInputFile(file_path, filename=file_name)
                    if file_ext == 'pdf':
                        await message.answer_document(file, caption=f"PDF документ: {file_name}")
                    elif file_ext == 'docx':
                        await message.answer_document(file, caption=f"Word документ: {file_name}")
                    elif file_ext == 'txt':
                        await message.answer_document(file, caption=f"Текстовый файл: {file_name}")
                    elif file_ext in ['mp4', 'avi', 'mov', 'mkv']:
                        await message.answer_video(file, caption=f"Видео файл: {file_name}")
                except Exception as e:
                    logging.error(f"Ошибка при отправке файла {file_path}: {e}")
                    await message.answer(f"Не удалось отправить файл {file_name}")

    except Exception as e:
        logging.error(f"Ошибка при обработке запроса: {e}")
        await message.answer(
            "Произошла ошибка при обработке запроса. Попробуйте переформулировать вопрос.",
            reply_markup=get_exit_button_ai()
        )


asyncio.run(initialize_vector_store())